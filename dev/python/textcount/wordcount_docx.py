#TODO: take care of encoding
#to make this compete with TextCount:
#TODO: add multiple formats (pdf, rtf, xls, etc.)
#TODO: add calculation of cost (setting of price per word or price per line)
#TODO: calculate number of lines from character count (~55 chars per line)

import docx

fname = 'A5.docx'

f = docx.Document(fname)

parac=len(f.paragraphs)
wc=0
cc=0
for p in f.paragraphs:
	cc += len(p.text)
	wc += len(p.text.split())

print 'paragraphs: %d' % (parac)
print 'words: %d' % (wc)
print 'chars: %d' % (cc)
